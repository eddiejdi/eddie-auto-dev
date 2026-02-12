#!/usr/bin/env python3
"""
Localiza o bloco de experi√™ncia existente para B3 e insere as linhas aprovadas
logo ap√≥s "Principais responsabilidades:" mantendo estilos quando poss√≠vel.
"""
import io
import json
import subprocess
import sys
import tempfile
from pathlib import Path

from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload, MediaFileUpload
from google.oauth2.credentials import Credentials
from google.auth.transport.requests import Request
from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

DOC_FILE_ID = "1y2eeV4No2zQD_ezeZCaBZiuswvANF8V3"
SECRETS_AGENT_HOST = "192.168.15.2"
SECRETS_AGENT_TOKEN_PATH = "/var/lib/eddie/secrets_agent/audit.db"
TOKEN_SECRET_NAME = "google/gdrive_token_edenilson_teixeira"
TOKEN_SECRET_FIELD = "token_json"

# Texto aprobado a inserir (cada item √© um bullet)
NEW_BULLETS = [
    "AIOps: implementa√ß√£o e opera√ß√£o de solu√ß√µes para detec√ß√£o pr√≥-ativa, alerting e automa√ß√£o de respostas a anomalias.",
    "Pipelines: cria√ß√£o e manuten√ß√£o de pipelines CI/CD e automa√ß√µes de deploy e rollback.",
    "Flyway: gest√£o e automa√ß√£o de migrations de banco de dados com Flyway.",
    "Atendimento de incidentes: participa√ß√£o em on-call, runbooks e postmortems para redu√ß√£o de MTTR.",
    "Modelos LLM: instala√ß√£o, configura√ß√£o e deploy de modelos LLM, integra√ß√£o e tuning.",
    "Integra√ß√£o: integra√ß√£o entre sistemas legados e novos servi√ßos via APIs e mensageria.",
    "Datalake: suporte a pipelines de ingest√£o (ETL/ELT) e organiza√ß√£o de dados para an√°lises.",
]

TARGET_HEADING_SNIPPET = "Cargo: Software Engineer"
TARGET_COMPANY_SNIPPET = "B3 (Bolsa Balc√£o do Brasil)"


def get_token_from_agent():
    cmd = [
        "ssh", f"homelab@{SECRETS_AGENT_HOST}",
        "python3 - <<'PY'\nimport sqlite3, base64, sys\nconn = sqlite3.connect('" + SECRETS_AGENT_TOKEN_PATH + "')\nc = conn.cursor()\nc.execute(\"SELECT value FROM secrets_store WHERE name='" + TOKEN_SECRET_NAME + "' AND field='" + TOKEN_SECRET_FIELD + "'\")\nrow = c.fetchone()\nconn.close()\nif not row:\n    print('NOT_FOUND')\n    sys.exit(0)\nprint(base64.b64decode(row[0]).decode())\nPY"
    ]
    res = subprocess.run(cmd, capture_output=True, text=True, timeout=20)
    if res.returncode != 0:
        raise RuntimeError(f"failed to fetch token: {res.stderr}")
    out = res.stdout.strip()
    if out == "NOT_FOUND":
        raise KeyError("token not found in secrets agent")
    return json.loads(out)


def insert_paragraph_after(paragraph, text=None, style=None):
    p = paragraph._p
    new_p = OxmlElement('w:p')
    p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    if style:
        try:
            new_para.style = style
        except Exception:
            pass
    return new_para


def main():
    print("üîê Obtendo token do Secrets Agent...")
    token_data = get_token_from_agent()
    creds = Credentials(
        token=token_data['token'],
        refresh_token=token_data.get('refresh_token'),
        token_uri=token_data.get('token_uri'),
        client_id=token_data.get('client_id'),
        client_secret=token_data.get('client_secret'),
        scopes=token_data.get('scopes', ['https://www.googleapis.com/auth/drive', 'https://www.googleapis.com/auth/documents']),
    )
    if creds.expired or not creds.valid:
        creds.refresh(Request())

    drive = build('drive', 'v3', credentials=creds)

    # baixar arquivo
    fh = io.BytesIO()
    request = drive.files().get_media(fileId=DOC_FILE_ID)
    downloader = MediaIoBaseDownload(fh, request)
    done = False
    print('üì• Baixando arquivo original (.docx)...')
    while not done:
        status, done = downloader.next_chunk()
    fh.seek(0)

    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tf:
        tf.write(fh.read())
        tmp_name = tf.name
    print('  ‚Üí salvo em', tmp_name)

    doc = Document(tmp_name)

    # localizar bloco alvo
    paras = doc.paragraphs
    target_idx = None
    for i, p in enumerate(paras):
        txt = p.text.strip()
        if TARGET_HEADING_SNIPPET in txt and TARGET_COMPANY_SNIPPET in txt:
            target_idx = i
            break
    if target_idx is None:
        # tentar localizar apenas por Cargo snippet
        for i, p in enumerate(paras):
            if TARGET_HEADING_SNIPPET in p.text:
                target_idx = i
                break
    if target_idx is None:
        print('‚ùå Bloco alvo n√£o encontrado. Abortando.')
        sys.exit(1)

    print('üîé Bloco encontrado na posi√ß√£o', target_idx)

    # localizar "Principais responsabilidades:" dentro do bloco
    # consideramos bloco at√© pr√≥xima linha que come√ßa com "Cargo:" (exceto o pr√≥prio)
    end_idx = None
    for j in range(target_idx+1, len(paras)):
        if paras[j].text.strip().startswith('Cargo:'):
            end_idx = j
            break
    if end_idx is None:
        end_idx = len(paras)

    # buscar a linha "Principais responsabilidades:" entre target_idx and end_idx
    pr_idx = None
    for j in range(target_idx, end_idx):
        if paras[j].text.strip().lower().startswith('principais responsabilidades'):
            pr_idx = j
            break

    if pr_idx is None:
        # se nao existe, inserir um heading "Principais responsabilidades:" logo ap√≥s o cargo
        insert_at = paras[target_idx]
        title_para = insert_paragraph_after(insert_at, 'Principais responsabilidades:')
        pr_idx = paras.index(title_para)
        print('  ‚ÑπÔ∏è Criada se√ß√£o "Principais responsabilidades:" na posi√ß√£o', pr_idx)
    else:
        print('  ‚ÑπÔ∏è Encontrada se√ß√£o "Principais responsabilidades:" na posi√ß√£o', pr_idx)

    # Inserir bullets ap√≥s pr_idx
    # tentar detectar estilo de bullet existente
    bullet_style = None
    # procurar pr√≥ximo paragrafo com style 'List Bullet'
    for k in range(pr_idx+1, min(pr_idx+6, len(paras))):
        s = paras[k].style.name if paras[k].style is not None else ''
        if 'List' in s or 'Bullet' in s or 'Listing' in s:
            bullet_style = paras[k].style
            break
    if not bullet_style:
        # fallback para 'List Bullet' se existir
        try:
            bullet_style = doc.styles['List Bullet']
        except Exception:
            bullet_style = None

    # inserir em ordem
    insert_point = paras[pr_idx]
    for b in NEW_BULLETS:
        newp = insert_paragraph_after(insert_point, b, style=bullet_style)
        # next insert after the newly created paragraph
        insert_point = newp

    # salvar e substituir no Drive
    out_tmp = tmp_name + '.updated.docx'
    doc.save(out_tmp)
    print('üìù Atualiza√ß√£o salva em', out_tmp)

    media = MediaFileUpload(out_tmp, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document', resumable=True)
    print('üì§ Substituindo arquivo original no Drive...')
    updated = drive.files().update(fileId=DOC_FILE_ID, media_body=media, fields='id,webViewLink').execute()
    print('  ‚úÖ Substitu√≠do. Link:', updated.get('webViewLink'))

    # cleanup
    try:
        Path(tmp_name).unlink()
        Path(out_tmp).unlink()
    except Exception:
        pass


if __name__ == '__main__':
    main()
