#!/usr/bin/env python3
"""
Baixa o .docx original do Drive, insere a seção aprovada mantendo formatação
(e.g., adicionando heading + parágrafos), e atualiza o arquivo original no Drive.
Usa o token armazenado no Secrets Agent (homelab).
"""
import io
import json
import subprocess
import sys
import tempfile
from pathlib import Path

from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload, MediaFileUpload
from google.oauth2.credentials import Credentials
from google.auth.transport.requests import Request
from docx import Document

DOC_FILE_ID = "1y2eeV4No2zQD_ezeZCaBZiuswvANF8V3"
SECRETS_AGENT_HOST = "192.168.15.2"
SECRETS_AGENT_TOKEN_PATH = "/var/lib/shared/secrets_agent/audit.db"
TOKEN_SECRET_NAME = "google/gdrive_token_edenilson_teixeira"
TOKEN_SECRET_FIELD = "token_json"

INSERT_HEADING = "B3 S.A. — Analista de Operações / SRE (Jul/2024 — Presente)"
INSERT_LINES = [
    "Atuação em operações e confiabilidade de serviços com foco em automação, observabilidade e integração de plataformas críticas.",
    "AIOps & Observabilidade: Implementação e evolução de soluções AIOps para detecção pró‑ativa, alerting e automação de respostas a anomalias.",
    "Pipelines & CI/CD: Projetos e operação de pipelines de entrega contínua (build/deploy), incluindo automações para testes, rollbacks e versionamento.",
    "Banco de dados / Migrações: Gestão e automação de migrations com Flyway, garantindo deploys seguros e rastreáveis do schema.",
    "Atendimento de incidentes: Participação em on‑call, runbooks, resposta a incidentes e condução de postmortems para redução de MTTR.",
    "Modelos LLM: Instalação, configuração e operação de modelos LLM em produção (deploy, serving, tuneamento e integração com fluxos existentes).",
    "Integração de sistemas legados: Projetos de integração entre sistemas legados e microserviços via APIs e barramento de mensagens, com foco em compatibilidade e resiliência.",
    "Datalake & Ingestão: Concepção e suporte ao pipeline de ingestão para datalake (ETL/ELT), organização de dados e governança mínima para análises.",
]


def get_token_from_agent():
    cmd = [
        "ssh", f"homelab@{SECRETS_AGENT_HOST}",
        "python3 - <<'PY'\nimport sqlite3, base64, sys\nconn = sqlite3.connect('" + SECRETS_AGENT_TOKEN_PATH + "')\nc = conn.cursor()\nc.execute(\"SELECT value FROM secrets_store WHERE name='" + TOKEN_SECRET_NAME + "' AND field='" + TOKEN_SECRET_FIELD + "'\")\nrow = c.fetchone()\nconn.close()\nif not row:\n    print('NOT_FOUND')\n    sys.exit(0)\nprint(base64.b64decode(row[0]).decode())\nPY"
    ]
    res = subprocess.run(cmd, capture_output=True, text=True, timeout=20)
    if res.returncode != 0:
        raise RuntimeError(f"failed to fetch token: {res.stderr}")
    out = res.stdout.strip()
    if out == "NOT_FOUND":
        raise KeyError("token not found in secrets agent")
    return json.loads(out)


def main():
    print("🔐 Obtendo token do Secrets Agent...")
    token_data = get_token_from_agent()
    creds = Credentials(
        token=token_data['token'],
        refresh_token=token_data.get('refresh_token'),
        token_uri=token_data.get('token_uri'),
        client_id=token_data.get('client_id'),
        client_secret=token_data.get('client_secret'),
        scopes=token_data.get('scopes', ['https://www.googleapis.com/auth/drive']),
    )
    if creds.expired or not creds.valid:
        print('🔄 Renovando token...')
        creds.refresh(Request())
    drive = build('drive', 'v3', credentials=creds)

    meta = drive.files().get(fileId=DOC_FILE_ID, fields='name,mimeType').execute()
    name = meta.get('name')
    mime = meta.get('mimeType')
    print(f"Arquivo: {name} (mime: {mime})")

    # 1) download do .docx (ou do arquivo original)
    fh = io.BytesIO()
    request = drive.files().get_media(fileId=DOC_FILE_ID)
    downloader = MediaIoBaseDownload(fh, request)
    done = False
    print('📥 Baixando arquivo original...')
    while not done:
        status, done = downloader.next_chunk()
    fh.seek(0)

    # salvar temporário
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tf:
        tf.write(fh.read())
        tmp_name = tf.name
    print('  → salvo em', tmp_name)

    # 2) abrir com python-docx e inserir seção ao final preservando estilos
    doc = Document(tmp_name)
    # inserir heading com estilo de título do documento, se existir
    try:
        doc.add_heading(INSERT_HEADING, level=2)
    except Exception:
        doc.add_paragraph(INSERT_HEADING)
    for l in INSERT_LINES:
        p = doc.add_paragraph(l)
        p.style = doc.styles['Normal'] if 'Normal' in doc.styles else None
    # salvar modificado
    out_tmp = tmp_name + '.updated.docx'
    doc.save(out_tmp)
    print('📝 Atualização salva em', out_tmp)

    # 3) upload substituindo o arquivo original
    media = MediaFileUpload(out_tmp, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document', resumable=True)
    print('📤 Substituindo arquivo original no Drive...')
    updated = drive.files().update(fileId=DOC_FILE_ID, media_body=media, fields='id,webViewLink').execute()
    print('  ✅ Substituído. Link:', updated.get('webViewLink'))

    # cleanup
    try:
        Path(tmp_name).unlink()
        Path(out_tmp).unlink()
    except Exception:
        pass


if __name__ == '__main__':
    main()
